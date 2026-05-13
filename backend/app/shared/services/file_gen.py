# gerador de arquivos com logo da Tati
from __future__ import annotations

import io
import os
import re
import tempfile
import uuid
from datetime import datetime
from pathlib import Path

# Diretório temporário para arquivos gerados (limpo periodicamente)
FILES_DIR = Path(tempfile.gettempdir()) / "tati_files"
FILES_DIR.mkdir(exist_ok=True)

# caminho da logo
_LOGO_CANDIDATES = [
    Path(__file__).parent.parent / "frontend" / "assets" / "images" / "tati_logo.jpg",
    Path(__file__).parent.parent / "assets" / "images" / "tati_logo.jpg",
]
LOGO_PATH = next((p for p in _LOGO_CANDIDATES if p.exists()), None)

def _new_id() -> str:
    return uuid.uuid4().hex

def detect_file_request(text: str) -> dict | None:
    """
    Detectará se a resposta da IA contém pedido/geração de arquivo.
    Retorna dict com {format, title, content} ou None.
    
    A IA deve incluir um bloco especial na resposta:
    <<<FILE:pdf:Título do Arquivo>>>
    conteúdo do arquivo aqui
    <<<END_FILE>>>
    """
    pattern = re.compile(
        r'<<<FILE:(pdf|word|excel):([^>]+)>>>\n(.*?)<<<END_FILE>>>',
        re.DOTALL | re.IGNORECASE
    )
    match = pattern.search(text)
    if not match: return None
    return {
        "format": match.group(1).lower(),
        "title": match.group(2).strip(),
        "content": match.group(3).strip(),
        "full_match": match.group(0),
    }

def clean_response_text(text: str) -> str:
    """Remove o bloco FILE da resposta antes de exibir ao usuário."""
    pattern = re.compile(r'<<<FILE:.*?<<<END_FILE>>>', re.DOTALL | re.IGNORECASE)
    return pattern.sub('', text).strip()
 
 
def generate_file(fmt: str, title: str, content: str) -> tuple[str, str, str]:
    """
    Gera o arquivo e salva em FILES_DIR.
    Retorna (file_id, filename, filepath).
    """
    file_id = _new_id()
    if fmt == "pdf":
        filename, path = _gen_pdf(file_id, title, content)
    elif fmt == "word":
        filename, path = _gen_word(file_id, title, content)
    elif fmt == "excel":
        filename, path = _gen_excel(file_id, title, content)
    else:
        raise ValueError(f"Formato desconhecido: {fmt}")
    return file_id, filename, str(path)
 
 
# ── PDF ───────────────────────────────────────────────────────────────────────
 
def _gen_pdf(file_id: str, title: str, content: str) -> tuple[str, Path]:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
 
    filename = f"{_slugify(title)}.pdf"
    path = FILES_DIR / f"{file_id}_{filename}"
 
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm,
        topMargin=2*cm, bottomMargin=2*cm,
    )
 
    styles = getSampleStyleSheet()
    primary = colors.HexColor("#7c3aed")  # roxo da Tati
 
    style_title = ParagraphStyle(
        "TatiTitle", parent=styles["Title"],
        textColor=primary, fontSize=20, spaceAfter=6,
        fontName="Helvetica-Bold",
    )
    style_subtitle = ParagraphStyle(
        "TatiSub", parent=styles["Normal"],
        textColor=colors.HexColor("#6b7280"), fontSize=9,
        spaceAfter=12, alignment=TA_CENTER,
    )
    style_body = ParagraphStyle(
        "TatiBody", parent=styles["Normal"],
        fontSize=11, leading=16, spaceAfter=8,
        fontName="Helvetica",
    )
    style_h2 = ParagraphStyle(
        "TatiH2", parent=styles["Heading2"],
        textColor=primary, fontSize=13, spaceBefore=12, spaceAfter=4,
        fontName="Helvetica-Bold",
    )
 
    story = []
 
    # Logo
    if LOGO_PATH:
        try:
            img = Image(str(LOGO_PATH), width=2.5*cm, height=2.5*cm)
            img.hAlign = "CENTER"
            story.append(img)
            story.append(Spacer(1, 0.3*cm))
        except Exception:
            pass
 
    story.append(Paragraph(title, style_title))
    story.append(Paragraph(
        f"Teacher Tati AI · {datetime.now().strftime('%d/%m/%Y')}",
        style_subtitle,
    ))
    story.append(HRFlowable(width="100%", thickness=1, color=primary, spaceAfter=12))
 
    # Conteúdo — parse básico de markdown
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            story.append(Spacer(1, 0.2*cm))
            continue
        if line.startswith("## "):
            story.append(Paragraph(line[3:], style_h2))
        elif line.startswith("# "):
            story.append(Paragraph(line[2:], style_h2))
        elif line.startswith("- ") or line.startswith("• "):
            story.append(Paragraph(f"• {line[2:]}", style_body))
        elif re.match(r'^\d+\.', line):
            story.append(Paragraph(line, style_body))
        else:
            # Bold **texto**
            line = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', line)
            story.append(Paragraph(line, style_body))
 
    # Rodapé
    story.append(Spacer(1, 0.5*cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#e5e7eb")))
    story.append(Paragraph(
        "Gerado por Teacher Tati AI · tatiai.com.br",
        ParagraphStyle("footer", parent=styles["Normal"],
                       fontSize=8, textColor=colors.HexColor("#9ca3af"),
                       alignment=TA_CENTER, spaceBefore=4),
    ))
 
    doc.build(story)
    return filename, path
 
 
# ── Word ──────────────────────────────────────────────────────────────────────
 
def _gen_word(file_id: str, title: str, content: str) -> tuple[str, Path]:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
 
    filename = f"{_slugify(title)}.docx"
    path = FILES_DIR / f"{file_id}_{filename}"
 
    doc = Document()
 
    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
 
    purple = RGBColor(0x7c, 0x3a, 0xed)
 
    # Logo
    if LOGO_PATH:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(LOGO_PATH), width=Cm(2.5))
        except Exception:
            pass
 
    # Título
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = purple
        run.font.size = Pt(20)
 
    # Subtítulo
    sub = doc.add_paragraph(f"Teacher Tati AI · {datetime.now().strftime('%d/%m/%Y')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
 
    doc.add_paragraph()  # espaço
 
    # Conteúdo
    for line in content.split("\n"):
        line_s = line.strip()
        if not line_s:
            doc.add_paragraph()
            continue
        if line_s.startswith("## ") or line_s.startswith("# "):
            h2 = doc.add_heading(line_s.lstrip("# "), level=2)
            for run in h2.runs:
                run.font.color.rgb = purple
        elif line_s.startswith("- ") or line_s.startswith("• "):
            p = doc.add_paragraph(line_s[2:], style="List Bullet")
        elif re.match(r'^\d+\.', line_s):
            p = doc.add_paragraph(line_s, style="List Number")
        else:
            p = doc.add_paragraph()
            # Bold **texto**
            parts = re.split(r'(\*\*.*?\*\*)', line_s)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
 
    # Rodapé
    doc.add_paragraph()
    footer_p = doc.add_paragraph("Gerado por Teacher Tati AI · tatiai.com.br")
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.runs[0].font.size = Pt(8)
    footer_p.runs[0].font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)
 
    doc.save(str(path))
    return filename, path
 
 
# ── Excel ─────────────────────────────────────────────────────────────────────
 
def _gen_excel(file_id: str, title: str, content: str) -> tuple[str, Path]:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.drawing.image import Image as XLImage
    from openpyxl.utils import get_column_letter
 
    filename = f"{_slugify(title)}.xlsx"
    path = FILES_DIR / f"{file_id}_{filename}"
 
    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]
 
    purple = "7C3AED"
    light_purple = "EDE9FE"
    gray = "6B7280"
 
    header_font = Font(name="Calibri", bold=True, color="FFFFFF", size=12)
    header_fill = PatternFill("solid", fgColor=purple)
    header_align = Alignment(horizontal="center", vertical="center")
 
    title_font = Font(name="Calibri", bold=True, color=purple, size=14)
    sub_font = Font(name="Calibri", color=gray, size=9)
 
    thin = Side(style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
 
    row = 1
 
    # Logo
    if LOGO_PATH:
        try:
            img = XLImage(str(LOGO_PATH))
            img.width = 60
            img.height = 60
            ws.add_image(img, "A1")
            ws.row_dimensions[1].height = 50
            row = 4
        except Exception:
            pass
 
    # Título
    ws.cell(row=row, column=1, value=title).font = title_font
    ws.merge_cells(f"A{row}:F{row}")
    ws.cell(row=row, column=1).alignment = Alignment(horizontal="center")
    row += 1
 
    ws.cell(row=row, column=1,
            value=f"Teacher Tati AI · {datetime.now().strftime('%d/%m/%Y')}").font = sub_font
    ws.merge_cells(f"A{row}:F{row}")
    ws.cell(row=row, column=1).alignment = Alignment(horizontal="center")
    row += 2
 
    # Conteúdo — detecta tabelas markdown e listas
    lines = [l for l in content.split("\n")]
    i = 0
    while i < len(lines):
        line = lines[i].strip()
 
        if not line:
            row += 1
            i += 1
            continue
 
        # Tabela markdown: | col1 | col2 |
        if line.startswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            is_header = i + 1 < len(lines) and re.match(r'[\|\-\s]+', lines[i+1])
            for col_idx, cell_val in enumerate(cells, 1):
                cell = ws.cell(row=row, column=col_idx, value=cell_val)
                cell.border = border
                cell.alignment = Alignment(horizontal="center", wrap_text=True)
                if is_header:
                    cell.font = header_font
                    cell.fill = header_fill
                else:
                    if row % 2 == 0:
                        cell.fill = PatternFill("solid", fgColor=light_purple)
            row += 1
            if is_header:
                i += 2  # pula linha de separação
                continue
 
        elif line.startswith("## ") or line.startswith("# "):
            cell = ws.cell(row=row, column=1, value=line.lstrip("# "))
            cell.font = Font(name="Calibri", bold=True, color=purple, size=12)
            ws.merge_cells(f"A{row}:F{row}")
            row += 1
 
        elif line.startswith("- ") or line.startswith("• "):
            ws.cell(row=row, column=1, value="•")
            ws.cell(row=row, column=2, value=line[2:])
            row += 1
 
        elif re.match(r'^\d+\.', line):
            parts = line.split(".", 1)
            ws.cell(row=row, column=1, value=parts[0] + ".")
            ws.cell(row=row, column=2, value=parts[1].strip() if len(parts) > 1 else "")
            row += 1
 
        else:
            cell = ws.cell(row=row, column=1, value=re.sub(r'\*\*(.*?)\*\*', r'\1', line))
            ws.merge_cells(f"A{row}:F{row}")
            row += 1
 
        i += 1
 
    # Auto-width
    for col in ws.columns:
        max_len = 0
        col_letter = get_column_letter(col[0].column)
        for cell in col:
            try:
                if cell.value:
                    max_len = max(max_len, len(str(cell.value)))
            except Exception:
                pass
        ws.column_dimensions[col_letter].width = min(max_len + 4, 50)
 
    wb.save(str(path))
    return filename, path
 
 
# ── Helpers ───────────────────────────────────────────────────────────────────
 
def _slugify(text: str) -> str:
    text = text.lower()
    text = re.sub(r'[áàâã]', 'a', text)
    text = re.sub(r'[éê]', 'e', text)
    text = re.sub(r'[íî]', 'i', text)
    text = re.sub(r'[óôõ]', 'o', text)
    text = re.sub(r'[úû]', 'u', text)
    text = re.sub(r'ç', 'c', text)
    text = re.sub(r'[^a-z0-9]+', '_', text)
    return text.strip('_')[:50]
 
 
def cleanup_old_files(max_age_hours: int = 2) -> None:
    """Remove arquivos gerados há mais de max_age_hours horas."""
    import time
    now = time.time()
    for f in FILES_DIR.glob("*"):
        if now - f.stat().st_mtime > max_age_hours * 3600:
            f.unlink(missing_ok=True)